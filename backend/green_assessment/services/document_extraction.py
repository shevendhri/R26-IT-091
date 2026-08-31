import csv
import re
import shutil
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


MODULE_DIR = Path(__file__).resolve().parents[1]
EXTRACTED_TEXT_ROOT = MODULE_DIR / "extracted_text"
LOW_TEXT_THRESHOLD = 40
WINDOWS_TESSERACT_PATH = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")


class DocumentExtractionError(Exception):
    pass


@dataclass
class ExtractionResult:
    method: str
    text_path: str
    char_count: int


def extract_document_text(document, project_id: int) -> ExtractionResult:
    source_path = Path(document.storage_path)
    if not source_path.exists() or not source_path.is_file():
        raise DocumentExtractionError("Uploaded source file was not found.")

    file_type = document.file_type.lower()
    if file_type == "pdf":
        text, method = _extract_pdf_text(source_path)
    elif file_type == "docx":
        text, method = _extract_docx_text(source_path), "python-docx"
    elif file_type in {"xlsx", "xls"}:
        text, method = _extract_excel_text(source_path, file_type), "spreadsheet"
    elif file_type == "csv":
        text, method = _extract_csv_text(source_path), "csv"
    elif file_type in {"png", "jpg", "jpeg"}:
        text, method = _extract_image_text(source_path), "tesseract_ocr"
    else:
        raise DocumentExtractionError("Unsupported file type for extraction.")

    cleaned_text = clean_extracted_text(text)
    text_path = _write_extracted_text(project_id, document.id, cleaned_text)
    return ExtractionResult(
        method=method,
        text_path=str(text_path),
        char_count=len(cleaned_text),
    )


def read_extracted_text(document) -> str:
    if not document.extracted_text_path:
        return ""

    text_path = Path(document.extracted_text_path)
    if not text_path.exists() or not text_path.is_file():
        return ""

    return text_path.read_text(encoding="utf-8")


def clean_extracted_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\t", " ")
    text = re.sub(r"[ \f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _write_extracted_text(project_id: int, document_id: int, text: str) -> Path:
    output_dir = EXTRACTED_TEXT_ROOT / f"project_{project_id}"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"document_{document_id}.txt"
    output_path.write_text(text, encoding="utf-8")
    return output_path


def _extract_pdf_text(path: Path) -> tuple[str, str]:
    try:
        import pymupdf as fitz
    except ImportError as exc:
        raise DocumentExtractionError("PyMuPDF is not installed.") from exc

    page_texts = []
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            page_text = page.get_text("text")
            page_texts.append(_format_page_text(page_number, page_text))

    direct_text = "\n\n".join(page_texts)
    if len(clean_extracted_text(direct_text)) >= LOW_TEXT_THRESHOLD:
        return direct_text, "pymupdf_direct"

    return _extract_scanned_pdf_text(path), "pymupdf_render_tesseract_ocr"


def _extract_scanned_pdf_text(path: Path) -> str:
    try:
        import pymupdf as fitz
        from PIL import Image
    except ImportError as exc:
        raise DocumentExtractionError("PyMuPDF and Pillow are required for scanned PDF OCR.") from exc

    page_texts = []
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            page_text = _run_ocr(image)
            page_texts.append(_format_page_text(page_number, page_text))

    return "\n\n".join(page_texts)


def _format_page_text(page_number: int, text: str) -> str:
    return f"--- PAGE {page_number} ---\n{text or ''}"


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError("python-docx is not installed.") from exc

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"--- TABLE {table_index} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_excel_text(path: Path, file_type: str) -> str:
    if file_type == "xlsx":
        return _extract_xlsx_text(path)
    return _extract_xls_text(path)


def _extract_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentExtractionError("openpyxl is not installed.") from exc

    workbook = load_workbook(path, data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"--- SHEET: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def _extract_xls_text(path: Path) -> str:
    try:
        import pandas as pd
    except ImportError as exc:
        raise DocumentExtractionError("pandas is not installed.") from exc

    try:
        sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
    except Exception as exc:
        raise DocumentExtractionError(
            "Unable to read XLS file. Install xlrd if this file format is required."
        ) from exc

    parts = []
    for sheet_name, frame in sheets.items():
        parts.append(f"--- SHEET: {sheet_name} ---")
        for row in frame.fillna("").values.tolist():
            values = [str(value) for value in row if str(value).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _extract_csv_text(path: Path) -> str:
    encodings = ["utf-8-sig", "utf-8", "latin-1"]
    last_error = None
    for encoding in encodings:
        try:
            with path.open("r", encoding=encoding, newline="") as csv_file:
                reader = csv.reader(csv_file)
                return "\n".join(" | ".join(row) for row in reader)
        except UnicodeDecodeError as exc:
            last_error = exc
    raise DocumentExtractionError("Unable to decode CSV file.") from last_error


def _extract_image_text(path: Path) -> str:
    try:
        from PIL import Image
    except ImportError as exc:
        raise DocumentExtractionError("Pillow is not installed.") from exc

    with Image.open(path) as image:
        return _run_ocr(image)


def _run_ocr(image) -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise DocumentExtractionError("pytesseract is not installed.") from exc

    try:
        _configure_tesseract_command(pytesseract)
        return pytesseract.image_to_string(image)
    except Exception as exc:
        raise DocumentExtractionError(
            f"Local Tesseract OCR failed: {exc}"
        ) from exc


def _configure_tesseract_command(pytesseract_module) -> None:
    if shutil.which("tesseract"):
        pytesseract_module.pytesseract.tesseract_cmd = "tesseract"
        return

    if WINDOWS_TESSERACT_PATH.exists():
        pytesseract_module.pytesseract.tesseract_cmd = str(WINDOWS_TESSERACT_PATH)
